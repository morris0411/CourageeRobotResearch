from __future__ import annotations

from copy import deepcopy
import hashlib
import os
from pathlib import Path
import shutil
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "Manuscript_Edited_Clean.docx"
BACKUP_PATH = ROOT / "Manuscript_Edited_Clean_before_action_figure.docx"
TEMP_PATH = ROOT / "Manuscript_Edited_Clean.__tmp_action_figure.docx"
FALLBACK_PATH = ROOT / "Manuscript_Edited_Clean_with_action_figure.docx"
FIGURE_PATH = ROOT / "image" / "study2_conflict_action_followup.png"

FIGURE_CAPTION = (
    "Figure 7. Personal courage self-evaluation scores within the conflict conditions "
    "by preexisting courage tendency group and action. Both approach and avoidance "
    "motives were presented in both conditions; the conditions differed only in whether "
    "the robot performed the admonishing behavior."
)


def find_paragraph(document: Document, fragment: str) -> Paragraph:
    matches = [paragraph for paragraph in document.paragraphs if fragment in paragraph.text]
    if len(matches) != 1:
        raise ValueError(f"Expected one paragraph containing {fragment!r}; found {len(matches)}")
    return matches[0]


def copy_run_properties(source_run, target_run) -> None:
    if source_run is not None and source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    template_run = next((run for run in paragraph.runs if run.text), None)
    paragraph.clear()
    run = paragraph.add_run(text)
    copy_run_properties(template_run, run)


def replace_fragment(paragraph: Paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        raise ValueError(f"Text to replace was not found: {old!r}")
    set_paragraph_text(paragraph, paragraph.text.replace(old, new, 1))


def insert_paragraph_after(reference: Paragraph, text: str = "") -> Paragraph:
    new_element = OxmlElement("w:p")
    reference._p.addnext(new_element)
    paragraph = Paragraph(new_element, reference._parent)
    paragraph.style = reference.style
    if text:
        run = paragraph.add_run(text)
        template_run = next((item for item in reference.runs if item.text), None)
        copy_run_properties(template_run, run)
    return paragraph


def add_figure(document: Document, reference: Paragraph) -> Paragraph:
    figure_paragraph = insert_paragraph_after(reference)
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.space_before = Pt(6)
    figure_paragraph.paragraph_format.space_after = Pt(6)
    inline_shape = figure_paragraph.add_run().add_picture(
        str(FIGURE_PATH),
        width=Inches(6.4),
    )
    inline_shape._inline.docPr.set("name", "Figure 7")
    inline_shape._inline.docPr.set(
        "descr",
        "Interaction plot of post courage scores by pre-courage group and action "
        "within the conflict conditions.",
    )
    return figure_paragraph


def update_document(document: Document) -> None:
    count_paragraph = find_paragraph(document, "Number of figures and tables:")
    replace_fragment(
        count_paragraph,
        "Number of figures and tables: 6 figures and 5 tables",
        "Number of figures and tables: 7 figures and 5 tables",
    )

    abstract = find_paragraph(document, "Robots can externalize pre-action internal states")
    replace_fragment(
        abstract,
        "Admonishing behavior had no clear effect.",
        "Within the conflict presentations, admonishing behavior produced no detectable "
        "additional difference.",
    )

    analysis = find_paragraph(document, "In Study 2, post-stimulus personal courage")
    analysis_addition = (
        " Because motive direction differed between the two no-conflict conditions, the "
        "omnibus action effect did not isolate admonishing behavior from motive content. "
        "Therefore, to examine action while holding motive content constant, we conducted "
        "a follow-up two-way mixed analysis of variance restricted to the conflict "
        "conditions, with preexisting courage tendency group as a between-participant "
        "factor and action as a within-participant factor."
    )
    set_paragraph_text(analysis, analysis.text + analysis_addition)

    action_result = find_paragraph(
        document,
        "The presence or absence of action did not have a significant effect",
    )
    set_paragraph_text(
        action_result,
        "To isolate the effect of admonishing behavior from motive content, we conducted "
        "a follow-up analysis restricted to the conflict conditions, in which both "
        "approach and avoidance motives were presented. Neither the main effect of action, "
        "F(1, 124) = 0.061, p = 0.805, partial η² < 0.001, nor the interaction between "
        "preexisting courage tendency group and action, F(1, 124) = 0.061, p = 0.805, "
        "partial η² < 0.001, was significant. Thus, when both approach and avoidance "
        "motives were presented, admonishing behavior produced no detectable additional "
        "difference in observers’ self-evaluations of personal courage.",
    )

    figure_callout = find_paragraph(document, "These simple effects are shown in Figure 6")
    set_paragraph_text(
        figure_callout,
        "The conflict simple effects are shown in Figure 6, and the action comparison "
        "within the conflict conditions is shown in Figure 7.",
    )
    add_figure(document, figure_callout)

    discussion = find_paragraph(
        document,
        "In Study 2, the hypothesis that the low-courage group",
    )
    replace_fragment(
        discussion,
        "In addition, the presence or absence of action did not have a significant effect. "
        "Therefore, whether the robot performed the admonishing behavior did not appear to "
        "determine observers’ self-evaluations of personal courage.",
        "In the follow-up analysis restricted to the conflict conditions, neither the main "
        "effect of action nor its interaction with preexisting courage tendency group was "
        "significant. Thus, when both approach and avoidance motives were presented, whether "
        "the robot performed the admonishing behavior did not produce a detectable additional "
        "difference in observers’ self-evaluations of personal courage.",
    )

    figure_6_caption = find_paragraph(document, "Figure 6. Simple effects")
    insert_paragraph_after(figure_6_caption, FIGURE_CAPTION)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def validate_document(path: Path, original_shape_count: int) -> None:
    with ZipFile(path) as archive:
        bad_entry = archive.testzip()
        if bad_entry is not None:
            raise ValueError(f"Invalid DOCX ZIP entry: {bad_entry}")
        figure_hash = sha256(FIGURE_PATH)
        media_hashes = {
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
        if figure_hash not in media_hashes:
            raise ValueError("The generated figure was not embedded in the DOCX package.")

    check = Document(path)
    if len(check.inline_shapes) != original_shape_count + 1:
        raise ValueError("Unexpected number of inline figures after DOCX update.")

    all_text = "\n".join(paragraph.text for paragraph in check.paragraphs)
    required = [
        "Number of figures and tables: 7 figures and 5 tables",
        "follow-up two-way mixed analysis of variance restricted to the conflict conditions",
        "F(1, 124) = 0.061, p = 0.805",
        FIGURE_CAPTION,
    ]
    for text in required:
        if text not in all_text:
            raise ValueError(f"Required manuscript text is missing after save: {text!r}")

    prohibited = [
        "Admonishing behavior had no clear effect.",
        "The presence or absence of action did not have a significant effect",
    ]
    for text in prohibited:
        if text in all_text:
            raise ValueError(f"Outdated broad claim remains after save: {text!r}")


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not FIGURE_PATH.exists():
        raise FileNotFoundError(FIGURE_PATH)

    document = Document(DOCX_PATH)
    if any(paragraph.text.startswith("Figure 7.") for paragraph in document.paragraphs):
        raise ValueError("Figure 7 is already present; refusing to insert a duplicate.")

    original_shape_count = len(document.inline_shapes)
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    update_document(document)
    if TEMP_PATH.exists():
        TEMP_PATH.unlink()
    document.save(TEMP_PATH)
    validate_document(TEMP_PATH, original_shape_count)

    try:
        os.replace(TEMP_PATH, DOCX_PATH)
        destination = DOCX_PATH
    except PermissionError:
        if FALLBACK_PATH.exists():
            FALLBACK_PATH.unlink()
        os.replace(TEMP_PATH, FALLBACK_PATH)
        destination = FALLBACK_PATH

    print(f"Updated manuscript: {destination}")
    print(f"Backup: {BACKUP_PATH}")
    print(f"Embedded figure: {FIGURE_PATH}")


if __name__ == "__main__":
    main()
