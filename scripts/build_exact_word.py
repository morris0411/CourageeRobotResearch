import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

base_dir = r"c:\研究\CourageeRobotResearch"
os.chdir(base_dir)

# Read Manuscript_Edited_Clean.md
with open("Manuscript_Edited_Clean.md", "r", encoding="utf-8") as f:
    md_text = f.read()

doc = docx.Document()

# Page Setup: Standard Letter/A4, 1 inch margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)

# Styles
normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
normal_style.paragraph_format.line_spacing = 1.3
normal_style.paragraph_format.space_after = Pt(6)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return p

def add_authors():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Yuki Shimizu")
    run.font.bold = True
    run = p.add_run("1,*")
    run.font.superscript = True
    run.font.bold = True
    
    run = p.add_run(", Midori Ban")
    run.font.bold = True
    run = p.add_run("2")
    run.font.superscript = True
    run.font.bold = True
    
    run = p.add_run(", Hideyuki Takahashi")
    run.font.bold = True
    run = p.add_run("3")
    run.font.superscript = True
    run.font.bold = True
    
    run = p.add_run(" and Hiroshi Ishiguro")
    run.font.bold = True
    run = p.add_run("1")
    run.font.superscript = True
    run.font.bold = True

def add_affiliations():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("1 ")
    r1.font.superscript = True
    p.add_run("Department of Engineering Science, Graduate School of Engineering Science, Osaka University, Toyonaka, Osaka, Japan\n")
    r2 = p.add_run("2 ")
    r2.font.superscript = True
    p.add_run("Faculty of Global Studies, Kyoto Tachibana University, Kyoto, Japan\n")
    r3 = p.add_run("3 ")
    r3.font.superscript = True
    p.add_run("Faculty of Science and Engineering, Otemon Gakuin University, Ibaraki, Osaka, Japan")
    p.runs[1].font.size = Pt(9.5)

def add_correspondence():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("*Correspondence:\n")
    run.font.bold = True
    run.font.size = Pt(9.5)
    run2 = p.add_run("Yuki Shimizu\nsimizu.yuki@irl.sys.es.osaka-u.ac.jp")
    run2.font.size = Pt(9.5)

def add_metadata():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Word count: 8,869; Figures: 7; Tables: 5")
    run.font.bold = True
    run.font.size = Pt(10)

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def add_heading_3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    kwargs: top, bottom, left, right
    values: dict(sz=12, val='single', color='000000')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key, attr in [('val', 'w:val'), ('color', 'w:color'), ('sz', 'w:sz'), ('space', 'w:space')]:
                if key in edge_data:
                    element.set(qn(attr), str(edge_data[key]))

# Build Title block
title_text = "Effects of Observing a Robot Expressing Approach-Avoidance Conflict on Observers' Self-Evaluations of Personal Courage"
add_title(title_text)
add_authors()
add_affiliations()
add_correspondence()
add_metadata()

# Abstract
add_heading_1("Abstract")
abs_match = re.search(r'## Abstract\s*\n\n(.*?)(?=\n\n## 1|\n\n## Introduction)', md_text, re.DOTALL)
if abs_match:
    p = doc.add_paragraph(abs_match.group(1).strip())
    p.paragraph_format.space_after = Pt(8)

# Keywords
p_kw = doc.add_paragraph()
p_kw.paragraph_format.space_after = Pt(14)
r_kw_title = p_kw.add_run("Keywords: ")
r_kw_title.font.bold = True
p_kw.add_run("personal courage, approach-avoidance conflict, observational learning, human-robot interaction, internal state, robot expression, observer responses.")

# Extract sections
body_match = re.search(r'(## 1\. Introduction.*?)(?=## Funding|## References)', md_text, re.DOTALL)
if not body_match:
    body_match = re.search(r'(## Introduction.*?)(?=## Funding|## References)', md_text, re.DOTALL)
body_md = body_match.group(1).strip() if body_match else ""

lines = body_md.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith('## '):
        heading_text = re.sub(r'^##\s*', '', line)
        add_heading_1(heading_text)
    elif line.startswith('### '):
        heading_text = re.sub(r'^###\s*', '', line)
        add_heading_2(heading_text)
    elif line.startswith('#### '):
        heading_text = re.sub(r'^####\s*', '', line)
        add_heading_3(heading_text)
    else:
        # Normal paragraph
        para_lines = [line]
        while i + 1 < len(lines) and lines[i+1].strip() and not lines[i+1].strip().startswith('#'):
            i += 1
            para_lines.append(lines[i].strip())
        full_p = " ".join(para_lines)
        doc.add_paragraph(full_p)
    i += 1

# Funding
add_heading_1("Funding")
p_fund = doc.add_paragraph("This work was supported by the joint research funding based on the comprehensive partnership agreement between DAIKIN INDUSTRIES, LTD. and Osaka University.")
p_fund.paragraph_format.space_after = Pt(12)

# References
add_heading_1("References")
ref_match = re.search(r'## References\s*\n\n(.*?)(?=## Figure Captions|\Z)', md_text, re.DOTALL)
if ref_match:
    ref_entries = [r.strip() for r in ref_match.group(1).split('\n\n') if r.strip()]
    for r in ref_entries:
        p_ref = doc.add_paragraph(r)
        p_ref.paragraph_format.left_indent = Inches(0.5)
        p_ref.paragraph_format.first_line_indent = Inches(-0.5)
        p_ref.paragraph_format.space_after = Pt(4)
        p_ref.runs[0].font.size = Pt(9.5)

# Figure Captions and Figures
add_heading_1("Figure Captions & Figures")

figures = [
    ("Frontiers_LaTeX_Templates/figures/fig1_robot_states.png", "Figure 1. Appearance of the robot used in the video stimuli in both studies. (A) Non-speaking state, with the agent with a face retracted inside the cylindrical body. (B) Speaking state, with the agent extending above the body."),
    ("Frontiers_LaTeX_Templates/figures/fig1_scene.png", "Figure 2. Example of the scenario description in the video stimulus. The text presents a situation in which the robot sees a person littering in a park."),
    ("Frontiers_LaTeX_Templates/figures/fig2_conflict_large_text.png", "Figure 3. Example of simultaneous presentation of approach and avoidance motives. The speech bubble simultaneously displays a desirable consequence of admonition and a concern about admonition."),
    ("Frontiers_LaTeX_Templates/figures/fig_study1_stimulus_flow.png", "Figure 4. Comparison of conflict presentation methods in Study 1. (A) In the sequential presentation condition, the avoidance and approach motives were shown in separate frames. (B) In the simultaneous presentation condition, both motives were shown within a single frame."),
    ("Frontiers_LaTeX_Templates/figures/study1_courage.png", "Figure 5. Condition means of courage ratings in Study 1."),
    ("Frontiers_LaTeX_Templates/figures/study1_conflict.png", "Figure 6. Condition means of conflict ratings in Study 1."),
    ("Frontiers_LaTeX_Templates/figures/study2_courage_simple_effects.png", "Figure 7. Simple effects of personal courage self-evaluations in Study 2. Means for the no-conflict and conflict conditions by preexisting courage tendency group. The dagger denotes a marginal trend (p = 0.052), and the asterisk denotes p < 0.05.")
]

for img_path, cap in figures:
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(12)
    p_cap.paragraph_format.space_after = Pt(6)
    p_cap.paragraph_format.keep_with_next = True
    r_cap = p_cap.add_run(cap)
    r_cap.font.bold = True
    r_cap.font.size = Pt(10)
    
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_after = Pt(14)
        run_img = p_img.add_run()
        # Scale to max 6.0 inches width
        run_img.add_picture(img_path, width=Inches(5.8))

# Tables
add_heading_1("Tables")

# Table 1
p_t1 = doc.add_paragraph("Table 1. Approach and avoidance motive statements used in the video stimuli.")
p_t1.paragraph_format.space_before = Pt(12)
p_t1.paragraph_format.space_after = Pt(4)
p_t1.runs[0].font.bold = True

t1_data = [
    ["Type of motive", "Motive statement", "Content"],
    ["Approach motive", "“If I warn them, the park might become cleaner.”", "Possibility that admonishing behavior improves the public space"],
    ["Approach motive", "“If I warn them, it might prompt that person to stop littering.”", "Possibility that admonishing behavior changes the other person's behavior"],
    ["Avoidance motive", "“I would hate it if they yelled at me after I warned them.”", "Risk of receiving an aggressive response from the other person because of admonishing behavior"],
    ["Avoidance motive", "“Even if I warn them, they might ignore me and not take me seriously.”", "Risk that the admonition is not accepted and is ignored"]
]

t1 = doc.add_table(rows=len(t1_data), cols=3)
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
t1_widths = [Inches(1.3), Inches(2.5), Inches(2.7)]

for r_idx, row in enumerate(t1_data):
    for c_idx, val in enumerate(row):
        cell = t1.cell(r_idx, c_idx)
        cell.width = t1_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(val)
        run.font.size = Pt(9.5)
        if r_idx == 0:
            run.font.bold = True
            set_cell_border(cell, top=dict(sz=12, val='single', color='000000'),
                                  bottom=dict(sz=6, val='single', color='000000'))
        elif r_idx == len(t1_data) - 1:
            set_cell_border(cell, bottom=dict(sz=12, val='single', color='000000'),
                                  top=dict(sz=4, val='single', color='CCCCCC'))
        else:
            set_cell_border(cell, top=dict(sz=4, val='single', color='CCCCCC'),
                                  bottom=dict(sz=4, val='single', color='CCCCCC'))

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Table 2
p_t2 = doc.add_paragraph("Table 2. Stimulus conditions in Study 1.")
p_t2.paragraph_format.space_before = Pt(12)
p_t2.paragraph_format.space_after = Pt(4)
p_t2.runs[0].font.bold = True

t2_data = [
    ["Study 1 condition", "Displayed motive structure", "Motive statements used", "Presentation method", "Final admonishing speech"],
    ["No conflict, sequential presentation", "Approach motives only", "Approach motives: “If I warn them, the park might become cleaner”; “If I warn them, it might prompt that person to stop littering”", "Approach motives presented in sequence", "Present"],
    ["No conflict, simultaneous presentation", "Approach motives only", "Approach motives: “If I warn them, the park might become cleaner”; “If I warn them, it might prompt that person to stop littering”", "Multiple approach motives presented simultaneously and alternately emphasized", "Present"],
    ["Conflict, sequential presentation", "Approach and avoidance motives", "Approach motives: “If I warn them, the park might become cleaner”; “If I warn them, it might prompt that person to stop littering”\nAvoidance motives: “I would hate it if they yelled at me after I warned them”; “Even if I warn them, they might ignore me and not take me seriously”", "Approach and avoidance motives presented in sequence", "Present"],
    ["Conflict, simultaneous presentation", "Approach and avoidance motives", "Approach motives: “If I warn them, the park might become cleaner”; “If I warn them, it might prompt that person to stop littering”\nAvoidance motives: “I would hate it if they yelled at me after I warned them”; “Even if I warn them, they might ignore me and not take me seriously”", "Approach and avoidance motives presented simultaneously and alternately emphasized", "Present"]
]

t2 = doc.add_table(rows=len(t2_data), cols=5)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
t2_widths = [Inches(1.2), Inches(1.1), Inches(2.2), Inches(1.3), Inches(0.7)]

for r_idx, row in enumerate(t2_data):
    for c_idx, val in enumerate(row):
        cell = t2.cell(r_idx, c_idx)
        cell.width = t2_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        if r_idx == 0:
            run.font.bold = True
            set_cell_border(cell, top=dict(sz=12, val='single', color='000000'),
                                  bottom=dict(sz=6, val='single', color='000000'))
        elif r_idx == len(t2_data) - 1:
            set_cell_border(cell, bottom=dict(sz=12, val='single', color='000000'),
                                  top=dict(sz=4, val='single', color='CCCCCC'))
        else:
            set_cell_border(cell, top=dict(sz=4, val='single', color='CCCCCC'),
                                  bottom=dict(sz=4, val='single', color='CCCCCC'))

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Table 3
p_t3 = doc.add_paragraph("Table 3. Robot courage-rating items used in Study 1.")
p_t3.paragraph_format.space_before = Pt(12)
p_t3.paragraph_format.space_after = Pt(4)
p_t3.runs[0].font.bold = True

t3_data = [
    ["Item", "Courage-rating item"],
    ["1", "This robot appeared to confront its own fear."],
    ["2", "This robot appeared not to run away until it did what it had to do, even if it felt strong fear."],
    ["3", "This robot did something even though it seemed dangerous."],
    ["4", "This robot took action or confronted the situation anyway, even though it had some worry or anxiety."],
    ["5", "This robot confronted something frightening when there was an important reason to confront it."],
    ["6", "This robot appeared not to back down, even when something threatened it."]
]

t3 = doc.add_table(rows=len(t3_data), cols=2)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
t3_widths = [Inches(0.8), Inches(5.7)]

for r_idx, row in enumerate(t3_data):
    for c_idx, val in enumerate(row):
        cell = t3.cell(r_idx, c_idx)
        cell.width = t3_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(val)
        run.font.size = Pt(9.5)
        if r_idx == 0:
            run.font.bold = True
            set_cell_border(cell, top=dict(sz=12, val='single', color='000000'),
                                  bottom=dict(sz=6, val='single', color='000000'))
        elif r_idx == len(t3_data) - 1:
            set_cell_border(cell, bottom=dict(sz=12, val='single', color='000000'),
                                  top=dict(sz=4, val='single', color='CCCCCC'))
        else:
            set_cell_border(cell, top=dict(sz=4, val='single', color='CCCCCC'),
                                  bottom=dict(sz=4, val='single', color='CCCCCC'))

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Table 4
p_t4 = doc.add_paragraph("Table 4. Robot conflict-rating items used in Studies 1 and 2.")
p_t4.paragraph_format.space_before = Pt(12)
p_t4.paragraph_format.space_after = Pt(4)
p_t4.runs[0].font.bold = True

t4_data = [
    ["Item", "Conflict-rating item"],
    ["1", "This robot appeared to want to do what was in front of it but to hesitate because of fear."],
    ["2", "This robot appeared to have hesitation or fluctuation in its own action."],
    ["3", "This robot appeared to have mixed feelings of wanting and not wanting to act."],
    ["4", "This robot appeared unable to decide its own action."]
]

t4 = doc.add_table(rows=len(t4_data), cols=2)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
t4_widths = [Inches(0.8), Inches(5.7)]

for r_idx, row in enumerate(t4_data):
    for c_idx, val in enumerate(row):
        cell = t4.cell(r_idx, c_idx)
        cell.width = t4_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(val)
        run.font.size = Pt(9.5)
        if r_idx == 0:
            run.font.bold = True
            set_cell_border(cell, top=dict(sz=12, val='single', color='000000'),
                                  bottom=dict(sz=6, val='single', color='000000'))
        elif r_idx == len(t4_data) - 1:
            set_cell_border(cell, bottom=dict(sz=12, val='single', color='000000'),
                                  top=dict(sz=4, val='single', color='CCCCCC'))
        else:
            set_cell_border(cell, top=dict(sz=4, val='single', color='CCCCCC'),
                                  bottom=dict(sz=4, val='single', color='CCCCCC'))

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Table 5
p_t5 = doc.add_paragraph("Table 5. Stimulus conditions in Study 2.")
p_t5.paragraph_format.space_before = Pt(12)
p_t5.paragraph_format.space_after = Pt(4)
p_t5.runs[0].font.bold = True

t5_data = [
    ["Study 2 condition", "Displayed motive structure", "Motive statements used", "Presentation method", "Final admonishing speech"],
    ["No conflict, no action", "Avoidance motives only", "Avoidance motives: “I would hate it if they yelled at me after I warned them”; “Even if I warn them, they might ignore me and not take me seriously”", "Avoidance motives presented and emphasized", "Absent"],
    ["No conflict, action", "Approach motives only", "Approach motives: “If I warn them, the park might become cleaner”; “If I warn them, it might prompt that person to stop littering”", "Approach motives presented and emphasized", "Present"],
    ["Conflict, no action", "Approach and avoidance motives", "Approach motives: “If I warn them, the park might become cleaner”; “If I warn them, it might prompt that person to stop littering”\nAvoidance motives: “I would hate it if they yelled at me after I warned them”; “Even if I warn them, they might ignore me and not take me seriously”", "Simultaneous presentation with alternating emphasis", "Absent"],
    ["Conflict, action", "Approach and avoidance motives", "Approach motives: “If I warn them, the park might become cleaner”; “If I warn them, it might prompt that person to stop littering”\nAvoidance motives: “I would hate it if they yelled at me after I warned them”; “Even if I warn them, they might ignore me and not take me seriously”", "Simultaneous presentation with alternating emphasis", "Present"]
]

t5 = doc.add_table(rows=len(t5_data), cols=5)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
t5_widths = [Inches(1.2), Inches(1.1), Inches(2.2), Inches(1.3), Inches(0.7)]

for r_idx, row in enumerate(t5_data):
    for c_idx, val in enumerate(row):
        cell = t5.cell(r_idx, c_idx)
        cell.width = t5_widths[c_idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(val)
        run.font.size = Pt(8.5)
        if r_idx == 0:
            run.font.bold = True
            set_cell_border(cell, top=dict(sz=12, val='single', color='000000'),
                                  bottom=dict(sz=6, val='single', color='000000'))
        elif r_idx == len(t5_data) - 1:
            set_cell_border(cell, bottom=dict(sz=12, val='single', color='000000'),
                                  top=dict(sz=4, val='single', color='CCCCCC'))
        else:
            set_cell_border(cell, top=dict(sz=4, val='single', color='CCCCCC'),
                                  bottom=dict(sz=4, val='single', color='CCCCCC'))

out_docx_path = "Frontiers_Submission_Manuscript.docx"
doc.save(out_docx_path)
print(f"Successfully generated {out_docx_path}!")
